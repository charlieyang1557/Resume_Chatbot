"""Enhanced data loader with PDF/Doc support and intelligent chunking."""

from __future__ import annotations

import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Sequence, Dict, Any
import hashlib


@dataclass(frozen=True)
class EnhancedDocument:
    """Enhanced document container with chunking metadata."""
    
    content: str
    metadata: Dict[str, Any]
    chunk_id: str
    source_hash: str


class DocumentChunker:
    """Intelligent document chunking with overlap and metadata preservation."""
    
    def __init__(self, chunk_size: int = 1000, overlap: int = 200):
        self.chunk_size = chunk_size
        self.overlap = overlap
    
    def _calculate_hash(self, content: str) -> str:
        """Calculate hash for content deduplication."""
        return hashlib.md5(content.encode()).hexdigest()[:12]
    
    def _split_text_intelligently(self, text: str, title: str = "") -> List[Dict[str, Any]]:
        """Split text into chunks with intelligent boundaries."""
        chunks = []
        
        # Split by sentences first
        sentences = re.split(r'(?<=[.!?])\s+', text)
        if not sentences:
            return []
        
        current_chunk = ""
        current_length = 0
        chunk_idx = 0
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            # If adding this sentence would exceed chunk size, finalize current chunk
            if current_length + len(sentence) > self.chunk_size and current_chunk:
                chunks.append({
                    'content': current_chunk.strip(),
                    'chunk_index': chunk_idx,
                    'title': title,
                    'length': current_length
                })
                
                # Start new chunk with overlap
                overlap_text = self._get_overlap_text(current_chunk)
                current_chunk = overlap_text + " " + sentence
                current_length = len(current_chunk)
                chunk_idx += 1
            else:
                current_chunk += (" " + sentence) if current_chunk else sentence
                current_length = len(current_chunk)
        
        # Add final chunk
        if current_chunk.strip():
            chunks.append({
                'content': current_chunk.strip(),
                'chunk_index': chunk_idx,
                'title': title,
                'length': current_length
            })
        
        return chunks
    
    def _get_overlap_text(self, text: str) -> str:
        """Get overlap text from the end of previous chunk."""
        if len(text) <= self.overlap:
            return text
        return text[-self.overlap:]


class EnhancedDataLoader:
    """Enhanced data loader with PDF/Doc support."""
    
    def __init__(self, chunk_size: int = 1000, overlap: int = 200):
        self.chunker = DocumentChunker(chunk_size, overlap)
    
    def _read_text_file(self, path: Path) -> str:
        """Read text file with proper encoding detection."""
        try:
            return path.read_text(encoding='utf-8')
        except UnicodeDecodeError:
            try:
                return path.read_text(encoding='latin-1')
            except UnicodeDecodeError:
                return path.read_text(encoding='utf-8', errors='ignore')
    
    def _read_pdf_file(self, path: Path) -> str:
        """Read PDF file using PyMuPDF (fitz)."""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(path)
            text = ""
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()
            return text
        except ImportError:
            raise ImportError("PyMuPDF required for PDF support. Install with: pip install PyMuPDF")
    
    def _read_doc_file(self, path: Path) -> str:
        """Read DOC file using python-docx."""
        try:
            from docx import Document
            doc = Document(path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except ImportError:
            raise ImportError("python-docx required for DOC support. Install with: pip install python-docx")
    
    def _read_json_file(self, path: Path) -> str:
        """Read JSON file and convert to text."""
        try:
            with open(path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Convert JSON to readable text
            if isinstance(data, dict):
                text_parts = []
                for key, value in data.items():
                    if isinstance(value, (str, int, float)):
                        text_parts.append(f"{key}: {value}")
                    elif isinstance(value, list):
                        text_parts.append(f"{key}: {', '.join(map(str, value))}")
                return "\n".join(text_parts)
            else:
                return str(data)
        except Exception:
            return ""
    
    def _parse_file(self, path: Path) -> str:
        """Parse file based on extension."""
        suffix = path.suffix.lower()
        
        if suffix in {'.md', '.txt'}:
            return self._read_text_file(path)
        elif suffix == '.pdf':
            return self._read_pdf_file(path)
        elif suffix in {'.doc', '.docx'}:
            return self._read_doc_file(path)
        elif suffix == '.json':
            return self._read_json_file(path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
    
    def _split_markdown_sections(self, text: str) -> List[tuple[str, str]]:
        """Split Markdown into sections with titles."""
        sections = []
        current_title = "Overview"
        current_lines = []
        
        for line in text.splitlines():
            line = line.rstrip()
            if line.startswith("#"):
                if current_lines:
                    sections.append((current_title, "\n".join(current_lines)))
                    current_lines = []
                current_title = line.lstrip("# ").strip() or current_title
            else:
                current_lines.append(line)
        
        if current_lines:
            sections.append((current_title, "\n".join(current_lines)))
        
        return sections if sections else [("Overview", text)]
    
    def load_documents(self, directory: Path) -> List[EnhancedDocument]:
        """Load and chunk documents from directory."""
        directory = directory.expanduser().resolve()
        if not directory.exists():
            return []
        
        documents = []
        supported_extensions = {'.md', '.txt', '.pdf', '.doc', '.docx', '.json'}
        
        for path in sorted(directory.rglob("*")):
            if path.is_file() and path.suffix.lower() in supported_extensions:
                try:
                    content = self._parse_file(path)
                    if not content.strip():
                        continue
                    
                    # Handle different file types
                    if path.suffix.lower() in {'.md', '.txt'}:
                        sections = self._split_markdown_sections(content)
                    else:
                        sections = [("Content", content)]
                    
                    # Create chunks for each section
                    for section_title, section_content in sections:
                        if not section_content.strip():
                            continue
                        
                        chunks = self.chunker._split_text_intelligently(
                            section_content, section_title
                        )
                        
                        for chunk_data in chunks:
                            source_hash = self.chunker._calculate_hash(content)
                            chunk_id = f"{path.stem}_{section_title}_{chunk_data['chunk_index']}"
                            
                            document = EnhancedDocument(
                                content=chunk_data['content'],
                                metadata={
                                    "source": str(path.relative_to(directory)),
                                    "title": section_title,
                                    "chunk_index": chunk_data['chunk_index'],
                                    "file_type": path.suffix.lower(),
                                    "chunk_length": chunk_data['length'],
                                    "full_document_hash": source_hash
                                },
                                chunk_id=chunk_id,
                                source_hash=source_hash
                            )
                            documents.append(document)
                
                except Exception as e:
                    print(f"Warning: Could not parse {path}: {e}")
                    continue
        
        return documents
    
    def save_corpus_jsonl(self, documents: List[EnhancedDocument], output_path: Path):
        """Save documents to JSONL format for external processing."""
        output_path = output_path.expanduser().resolve()
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            for doc in documents:
                record = {
                    "id": doc.chunk_id,
                    "content": doc.content,
                    "metadata": doc.metadata,
                    "source_hash": doc.source_hash
                }
                f.write(json.dumps(record, ensure_ascii=False) + '\n')


def load_resume_documents(directory: Path, chunk_size: int = 1000, overlap: int = 200) -> List[EnhancedDocument]:
    """Load resume documents with enhanced chunking.
    
    Parameters
    ----------
    directory : Path
        Directory containing resume files
    chunk_size : int, default=1000
        Maximum characters per chunk
    overlap : int, default=200
        Character overlap between chunks
        
    Returns
    -------
    List[EnhancedDocument]
        Chunked documents with metadata
    """
    loader = EnhancedDataLoader(chunk_size, overlap)
    return loader.load_documents(directory)


# Backward compatibility
def load_resume_documents_legacy(directory: Path) -> List:
    """Legacy function for backward compatibility."""
    from .data_loader import load_resume_documents as legacy_loader
    legacy_docs = legacy_loader(directory)
    
    # Convert to new format
    enhanced_docs = []
    for i, doc in enumerate(legacy_docs):
        enhanced_doc = EnhancedDocument(
            content=doc.content,
            metadata=doc.metadata,
            chunk_id=f"legacy_{i}",
            source_hash="legacy"
        )
        enhanced_docs.append(enhanced_doc)
    
    return enhanced_docs
